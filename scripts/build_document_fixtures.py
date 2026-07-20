"""Generate the deterministic binary document fixtures.

    python scripts/build_document_fixtures.py

The text fixtures (`.md`, `.html`, `.csv`, `.yaml`, `.json`, `.sql`) are checked
in as plain source and are not produced here — they are readable in a diff, which
is worth more than uniformity.

The three BINARY fixtures cannot be. They are generated instead:

    sample-requirements.docx   python-docx
    sample-tests.xlsx          openpyxl
    sample-design.pdf          hand-written PDF bytes (no dependency)
    sample-scanned.pdf         hand-written, one page, NO text  -> needs-ocr
    sample-encrypted.pdf       hand-written, /Encrypt dictionary -> encrypted

WHY DETERMINISM MATTERS HERE
----------------------------
A fixture whose CONTENT changes between runs would change its content hash, and
the whole point of the Asset model is that an unchanged document produces no new
Revision. So every timestamp is fixed, every id is fixed, and no generator field
is allowed to carry "now".

The achievable guarantee differs by format. The PDFs are byte-identical
everywhere, because this file writes their bytes directly. A DOCX/XLSX is a
DEFLATE archive, and the compressed container bytes depend on the platform's
zlib build — so those are reproducible member-for-member, not byte-for-byte,
across operating systems. CI compares them accordingly. Storing them
uncompressed would buy true byte-identity at the cost of inflating the DOCX from
37 KB to 832 KB, which is not worth it for a fixture.

The PDFs are written by hand rather than with a PDF library: it keeps the
repository free of a PDF *writing* dependency (we only need to *read* PDFs), the
files stay under 2 KB, and the bytes are completely under our control — which is
what makes "this page has no text" and "this file is encrypted" reliable test
inputs rather than a library's changing idea of them.

ALL CONTENT IS INVENTED. The "NameCheck" service, its requirements and its test
cases are fictional; nothing here comes from any real or proprietary document.
"""
from __future__ import annotations

import datetime as _dt
import sys
import zlib
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
FIXTURE_DIR = REPO_ROOT / "fixtures" / "documents"

#: Every generated timestamp. Fixed so a regenerated fixture has identical
#: content (see the module docstring on what "identical" means per format).
EPOCH = _dt.datetime(2026, 1, 2, 3, 4, 5)


def normalize_zip(path: Path) -> None:
    """Rewrite an OOXML package with fixed member metadata.

    ``python-docx`` and ``openpyxl`` stamp each ZIP member with the current time,
    so two runs a second apart produce different bytes — and different bytes mean
    a different content hash, which would make the fixture look like a changed
    document every time it was regenerated. Rewriting with a fixed
    ``date_time``, fixed compression and a fixed member order makes the output
    reproducible.

    ``[Content_Types].xml`` is written first because the OPC specification
    expects it at the start of the package; everything else is sorted so the
    order cannot depend on dictionary iteration.

    ``docProps/core.xml`` gets one more fix-up: ``openpyxl`` overwrites
    ``dcterms:modified`` with the wall clock during ``save()``, ignoring the
    value set on the workbook, so the recorded modification time is rewritten to
    the fixed epoch here.
    """
    import re
    import zipfile

    with zipfile.ZipFile(path) as source:
        members = {info.filename: source.read(info.filename)
                   for info in source.infolist()}
    core = members.get("docProps/core.xml")
    if core is not None:
        stamp = EPOCH.strftime("%Y-%m-%dT%H:%M:%SZ").encode("utf-8")
        members["docProps/core.xml"] = re.sub(
            rb"(<dcterms:(?:created|modified)\b[^>]*>)[^<]*(</dcterms:)",
            lambda m: m.group(1) + stamp + m.group(2), core)
    order = sorted(members, key=lambda n: (n != "[Content_Types].xml", n))
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED,
                         compresslevel=6) as target:
        for name in order:
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, members[name])


#: A 1x1 transparent PNG, written out byte-for-byte so the DOCX fixture can
#: carry an embedded image without a checked-in binary asset or an image
#: dependency. Its only job is to exercise the "embedded images are recorded as
#: unsupported content, never OCR'd" path.
_TINY_PNG = bytes([
    0x89, 0x50, 0x4E, 0x47, 0x0D, 0x0A, 0x1A, 0x0A,
    0x00, 0x00, 0x00, 0x0D, 0x49, 0x48, 0x44, 0x52,
    0x00, 0x00, 0x00, 0x01, 0x00, 0x00, 0x00, 0x01,
    0x08, 0x06, 0x00, 0x00, 0x00, 0x1F, 0x15, 0xC4,
    0x89, 0x00, 0x00, 0x00, 0x0A, 0x49, 0x44, 0x41,
    0x54, 0x78, 0x9C, 0x63, 0x00, 0x01, 0x00, 0x00,
    0x05, 0x00, 0x01, 0x0D, 0x0A, 0x2D, 0xB4, 0x00,
    0x00, 0x00, 0x00, 0x49, 0x45, 0x4E, 0x44, 0xAE,
    0x42, 0x60, 0x82,
])


def build_docx(path: Path) -> None:
    import io as _io

    import docx
    from docx.shared import Emu

    document = docx.Document()
    core = document.core_properties
    core.author = "OpenMind Fixture Generator"
    core.title = "NameCheck Requirements"
    core.created = EPOCH
    core.modified = EPOCH
    core.revision = 3
    core.language = "en-GB"

    document.add_heading("NameCheck Requirements", level=1)
    document.add_paragraph(
        "This document specifies the NameCheck screening service. All content "
        "is invented for testing.")

    document.add_heading("4. Functional Requirements", level=2)
    document.add_paragraph(
        "The service screens submitted names against the watch list.")

    document.add_heading("4.3 NameCheck", level=3)
    document.add_paragraph(
        "REQ-NC-017: a manual review must time out after 30 minutes and the "
        "case must return to the queue.")
    document.add_paragraph(
        "REQ-NC-018: the client retries a failed screening three times before "
        "reporting an error.")
    document.add_paragraph("The endpoint is POST /name-check.",
                           style="List Bullet")
    document.add_paragraph("Configuration key: namecheck.review.timeout.minutes",
                           style="List Bullet")

    table = document.add_table(rows=1, cols=3)
    header = table.rows[0].cells
    header[0].text = "Requirement"
    header[1].text = "Title"
    header[2].text = "Owner"
    for req, title, owner in (
            ("REQ-NC-017", "Manual review timeout", "Operations"),
            ("REQ-NC-018", "Client retry policy", "Engineering"),
            ("REQ-NC-019", "Reject an empty name", "Engineering")):
        cells = table.add_row().cells
        cells[0].text = req
        cells[1].text = title
        cells[2].text = owner

    # An embedded image, so the "recorded as unsupported content, never OCR'd"
    # path has something real to report.
    document.add_paragraph("Figure 1 — screening flow (image).")
    document.add_picture(_io.BytesIO(_TINY_PNG), width=Emu(914400))

    document.save(str(path))


def build_xlsx(path: Path) -> None:
    import openpyxl

    book = openpyxl.Workbook()
    props = book.properties
    props.creator = "OpenMind Fixture Generator"
    props.title = "NameCheck Test Cases"
    props.created = EPOCH
    props.modified = EPOCH

    sheet = book.active
    sheet.title = "Test Cases"
    rows = [
        ("Case ID", "Requirement", "Scenario", "Expected", "Owner"),
        ("TC-001", "REQ-NC-017", "Review left idle", "Case returns to queue", "QA"),
        ("TC-002", "REQ-NC-018", "Screening fails twice", "Third attempt succeeds", "QA"),
        ("TC-003", "REQ-NC-019", "Empty name submitted", "Rejected with NC-100", "QA"),
    ]
    for row in rows:
        sheet.append(row)
    # A merged title banner, so merged-cell metadata has something to report.
    sheet.insert_rows(1)
    sheet["A1"] = "NameCheck acceptance tests"
    sheet.merge_cells("A1:E1")
    # A formula, stored and read AS A FORMULA and never evaluated.
    sheet["G2"] = "=COUNTA(A3:A6)"

    totals = book.create_sheet("Totals")
    totals["A1"] = "Total cases"
    totals["B1"] = "=COUNTA('Test Cases'!A3:A6)"

    hidden = book.create_sheet("Scratch")
    hidden["A1"] = "internal scratch values"
    hidden.sheet_state = "hidden"

    book.save(str(path))


# ---------------------------------------------------------------------------
# Minimal hand-written PDFs
# ---------------------------------------------------------------------------
def _pdf(objects: list, trailer_extra: str = "") -> bytes:
    """Assemble a PDF from 1-based object bodies, fixing up the xref offsets."""
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode("latin-1")
        out += body if isinstance(body, bytes) else body.encode("latin-1")
        out += b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    out += (f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R{trailer_extra} "
            f">>\nstartxref\n{xref_at}\n%%EOF\n").encode("latin-1")
    return bytes(out)


def _text_stream(lines: list, start_y: int = 760) -> bytes:
    parts = ["BT", "/F1 12 Tf", f"1 0 0 1 72 {start_y} Tm", "14 TL"]
    for line in lines:
        escaped = (line.replace("\\", r"\\").replace("(", r"\(")
                   .replace(")", r"\)"))
        parts.append(f"({escaped}) Tj T*")
    parts.append("ET")
    return "\n".join(parts).encode("latin-1")


def _page_objects(pages: list) -> list:
    """Build the object list for a PDF with one content stream per page."""
    page_count = len(pages)
    first_page_obj = 4                       # 1 catalog, 2 pages, 3 font
    kids = " ".join(f"{first_page_obj + i * 2} 0 R" for i in range(page_count))
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    for i, lines in enumerate(pages):
        page_obj = first_page_obj + i * 2
        content_obj = page_obj + 1
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> "
            f"/Contents {content_obj} 0 R >>")
        stream = _text_stream(lines) if lines else b""
        objects.append(b"<< /Length " + str(len(stream)).encode("latin-1")
                       + b" >>\nstream\n" + stream + b"\nendstream")
    return objects


def build_pdf(path: Path) -> None:
    """A two-page, text-bearing PDF with document metadata."""
    objects = _page_objects([
        ["NameCheck Design Note", "",
         "The screening service exposes POST /name-check.",
         "Requirement REQ-NC-017 governs the manual review timeout."],
        ["Page 2: Error codes", "",
         "NC-100 is returned when the submitted name is empty.",
         "Configuration key namecheck.review.timeout.minutes applies."],
    ])
    objects.append(
        "<< /Title (NameCheck Design Note) /Author (OpenMind Fixture Generator) "
        "/Producer (OpenMind fixtures) /CreationDate (D:20260102030405Z) >>")
    path.write_bytes(_pdf(objects, trailer_extra=f" /Info {len(objects)} 0 R"))


def build_scanned_pdf(path: Path) -> None:
    """One page, a drawn rectangle, and NO text operators at all.

    This is the image-only case: a correct parser extracts nothing from it, and
    must report `needs-ocr` rather than inventing content.
    """
    stream = b"0.9 0.9 0.9 rg\n72 600 468 120 re\nf\n"
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("latin-1") + b" >>\nstream\n"
        + stream + b"\nendstream",
    ]
    path.write_bytes(_pdf(objects))


def build_encrypted_pdf(path: Path) -> None:
    """A PDF whose trailer declares standard encryption.

    The body is not really enciphered — what is under test is that OpenMind
    reports `encrypted` and extracts nothing, rather than returning whatever
    bytes it could scrape. Producing genuinely enciphered content would need a
    PDF-writing dependency for no additional coverage.
    """
    stream = _text_stream(["This content is not readable without the password."])
    objects = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        "/Resources << /Font << /F1 6 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("latin-1") + b" >>\nstream\n"
        + stream + b"\nendstream",
        # /R 4 /V 4 with a non-empty /O and /U: pypdf must not be able to open
        # it with the empty user password.
        "<< /Filter /Standard /V 4 /R 4 /Length 128 /P -1340 "
        "/O <2A3B4C5D6E7F8091A2B3C4D5E6F708192A3B4C5D6E7F8091A2B3C4D5E6F70819> "
        "/U <1122334455667788990011223344556677889900112233445566778899001122> "
        "/CF << /StdCF << /CFM /AESV2 /Length 16 >> >> /StmF /StdCF "
        "/StrF /StdCF >>",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    path.write_bytes(_pdf(
        objects,
        trailer_extra=" /Encrypt 5 0 R /ID [<0102030405060708090A0B0C0D0E0F10> "
                      "<0102030405060708090A0B0C0D0E0F10>]"))


def main() -> int:
    FIXTURE_DIR.mkdir(parents=True, exist_ok=True)
    builders = [
        ("sample-requirements.docx", build_docx, True),
        ("sample-tests.xlsx", build_xlsx, True),
        ("sample-design.pdf", build_pdf, False),
        ("sample-scanned.pdf", build_scanned_pdf, False),
        ("sample-encrypted.pdf", build_encrypted_pdf, False),
    ]
    for name, builder, is_ooxml in builders:
        target = FIXTURE_DIR / name
        try:
            builder(target)
        except ImportError as exc:
            print(f"  SKIP {name}: {exc}", file=sys.stderr)
            continue
        if is_ooxml:
            normalize_zip(target)
        print(f"  wrote {name}  ({target.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())

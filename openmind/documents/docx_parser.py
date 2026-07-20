"""DOCX (Office Open XML wordprocessing) via ``python-docx``.

SAFETY BEFORE PARSING
---------------------
A DOCX is a ZIP of XML. Both halves are hostile-input surfaces, so both are
bounded before ``python-docx`` sees a single byte:

1. :func:`~openmind.documents.security.inspect_zip` validates the package —
   member count, per-member size, total expansion, compression ratio and path
   traversal. Nothing is ever extracted to a filesystem path; the archive is
   read from memory.
2. :func:`~openmind.documents.security.harden_xml` applies
   ``defusedxml.defuse_stdlib()`` so the stdlib XML parsers ``python-docx`` uses
   refuse external entities, external DTDs and entity-expansion bombs. If
   ``defusedxml`` is not installed this parser declines the document rather than
   parsing it unsafely — an unsafe parse is never the fallback.

Macros are never executed, external links are never followed, and embedded
images are recorded as unsupported content rather than silently dropped.

STRUCTURE
---------
Heading level comes from the paragraph's STYLE (``Heading 1``…``Heading 9``,
``Title``), which is what the author actually asserted — never from font size or
from prose that looks like a heading. Body order is recovered by walking the
document body's XML children, so a table between two paragraphs stays between
them instead of being appended after every paragraph, which is what iterating
``document.paragraphs`` then ``document.tables`` would produce.
"""
from __future__ import annotations

import io
import re
from typing import Any, Dict, Iterator, List, Optional, Tuple

from ..domain.types import ContentMode, DocumentBlockType, DocumentParseStatus
from .builder import DocumentBuilder, slug
from .models import (DocumentParseContext, DocumentProbe, ParsedDocument,
                     dependency_unavailable)
from .probe import _DOCX_MARKER
from .security import DocumentSecurityError, harden_xml, inspect_zip

_MEDIA = ("application/vnd.openxmlformats-officedocument"
          ".wordprocessingml.document")

_HEADING_STYLE_RE = re.compile(r"^heading\s*(\d)$", re.IGNORECASE)
_CODE_STYLE_RE = re.compile(r"(code|source|listing|monospace|preformatted)",
                            re.IGNORECASE)
_LIST_STYLE_RE = re.compile(r"(list|bullet)", re.IGNORECASE)

#: OOXML namespace for wordprocessing elements, used for the body walk.
_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def CLAIMS(probe: DocumentProbe) -> bool:      # noqa: N802 - registry protocol
    """A DOCX is claimed on PACKAGE STRUCTURE, never on the suffix.

    A renamed ZIP with a ``.docx`` extension does not contain
    ``word/document.xml`` and is therefore not claimed — it falls through to be
    reported as unsupported, which is the honest answer.
    """
    return probe.is_zip_package and probe.has_zip_member(_DOCX_MARKER)


class DocxParser:
    """Paragraphs, heading hierarchy, lists and tables from a DOCX package."""

    name = "docx"
    version = "1.0"

    def supports(self, probe: DocumentProbe) -> bool:
        return CLAIMS(probe)

    def parse(self, content: bytes,
              context: DocumentParseContext) -> ParsedDocument:
        limits = context.limits
        doc = ParsedDocument(parser_name=self.name, parser_version=self.version,
                             media_type=_MEDIA,
                             title=context.filename or context.logical_key)

        if not harden_xml():
            return dependency_unavailable(context.filename, self.name,
                                          "defusedxml")
        try:
            import docx                                    # python-docx
        except ImportError:
            return dependency_unavailable(context.filename, self.name,
                                          "python-docx")

        try:
            package = inspect_zip(
                content, max_members=limits.zip_max_members,
                max_total_bytes=limits.zip_max_total_bytes,
                max_member_bytes=limits.zip_max_member_bytes,
                max_ratio=limits.zip_max_ratio)
        except DocumentSecurityError as exc:
            doc.status = DocumentParseStatus.UNSUPPORTED
            doc.reason = exc.code
            doc.add_warning(exc.code, exc.message, **exc.detail)
            return doc

        images = sum(1 for m in package["members"] if m.startswith("word/media/"))
        if images:
            doc.note_unsupported(
                "embedded-image", images,
                "embedded images are recorded but not extracted; OCR is not "
                "performed in this phase")
        macros = sum(1 for m in package["members"]
                     if m.startswith("word/vbaProject"))
        if macros:
            doc.note_unsupported("macro", macros,
                                 "macros are never read or executed")

        document = docx.Document(io.BytesIO(content))
        _read_core_properties(document, doc)

        builder = DocumentBuilder(doc, max_blocks=limits.max_blocks,
                                  max_block_chars=limits.max_block_chars)
        key = context.logical_key
        builder.add_root(doc.title, {"kind": "docx-paragraph", "document": key,
                                     "paragraphIndex": 0})
        _walk_body(document, builder, key, doc, limits)

        for block in doc.blocks:
            if (block.block_type == DocumentBlockType.HEADING
                    and block.metadata.get("level", 9) <= 1):
                doc.title = block.text.strip() or doc.title
                doc.blocks[0].text = doc.title
                break
        return doc


def _read_core_properties(document: Any, doc: ParsedDocument) -> None:
    """Copy the OOXML core properties. Never raises — a document with a damaged
    or absent core-properties part is still perfectly parseable."""
    try:
        props = document.core_properties
    except Exception:
        return
    def _text(value: Any) -> str:
        if value is None:
            return ""
        return value.isoformat() if hasattr(value, "isoformat") else str(value)

    doc.metadata.author = _text(getattr(props, "author", ""))
    doc.metadata.created = _text(getattr(props, "created", None))
    doc.metadata.modified = _text(getattr(props, "modified", None))
    doc.metadata.language = _text(getattr(props, "language", ""))
    # `revision` is an explicit, documented DOCX field — the ONLY place a
    # version label may come from here. Nothing is inferred from prose.
    revision = getattr(props, "revision", None)
    if revision:
        doc.metadata.version_label = str(revision)
    title = _text(getattr(props, "title", ""))
    if title.strip():
        doc.title = title.strip()
    subject = _text(getattr(props, "subject", ""))
    if subject:
        doc.metadata.extra["subject"] = subject
    category = _text(getattr(props, "category", ""))
    if category:
        doc.metadata.extra["category"] = category


def _iter_body(document: Any) -> Iterator[Tuple[str, Any]]:
    """Yield ``("paragraph"|"table", element)`` in true document order.

    ``document.paragraphs`` and ``document.tables`` are two separate flat lists,
    so using them loses the interleaving — every table would land after every
    paragraph and a table's caption would end up nowhere near it. Walking the
    body's XML children preserves the order the author wrote.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for child in body.iterchildren():
        if child.tag == f"{_W}p":
            yield "paragraph", Paragraph(child, document)
        elif child.tag == f"{_W}tbl":
            yield "table", Table(child, document)


def _style_name(paragraph: Any) -> str:
    try:
        return (paragraph.style.name or "").strip()
    except Exception:
        return ""


def _heading_level(style: str) -> Optional[int]:
    if style.lower() == "title":
        return 1
    match = _HEADING_STYLE_RE.match(style)
    return int(match.group(1)) if match else None


def _column_letter(index: int) -> str:
    letters = ""
    while index > 0:
        index, rem = divmod(index - 1, 26)
        letters = chr(ord("A") + rem) + letters
    return letters or "A"


def _walk_body(document: Any, builder: DocumentBuilder, key: str,
               doc: ParsedDocument, limits: Any) -> None:
    paragraph_index = 0
    table_index = 0
    truncated_paragraphs = False
    truncated_tables = False

    for kind, element in _iter_body(document):
        if builder.full:
            break
        if kind == "paragraph":
            paragraph_index += 1
            if paragraph_index > limits.docx_max_paragraphs:
                if not truncated_paragraphs:
                    truncated_paragraphs = True
                    doc.mark_partial(
                        "limit-exceeded",
                        f"the document has more than DOCX_MAX_PARAGRAPHS "
                        f"({limits.docx_max_paragraphs}) paragraphs; the "
                        f"remainder was not extracted",
                        limit="DOCX_MAX_PARAGRAPHS",
                        allowed=limits.docx_max_paragraphs)
                continue
            _emit_paragraph(builder, key, element, paragraph_index)
        else:
            table_index += 1
            if table_index > limits.docx_max_tables:
                if not truncated_tables:
                    truncated_tables = True
                    doc.mark_partial(
                        "limit-exceeded",
                        f"the document has more than DOCX_MAX_TABLES "
                        f"({limits.docx_max_tables}) tables; the remainder was "
                        f"not extracted",
                        limit="DOCX_MAX_TABLES", allowed=limits.docx_max_tables)
                continue
            _emit_table(builder, key, element, table_index)

    doc.coverage["paragraphs"] = paragraph_index
    doc.coverage["tables"] = table_index


def _emit_paragraph(builder: DocumentBuilder, key: str, paragraph: Any,
                    index: int) -> None:
    text = (paragraph.text or "").strip()
    if not text:
        return
    style = _style_name(paragraph)
    locator: Dict[str, Any] = {"kind": "docx-paragraph", "document": key,
                               "paragraphIndex": index}
    level = _heading_level(style)
    if level is not None:
        locator["headingPath"] = list(builder.heading_path)
        block = builder.add(
            DocumentBlockType.HEADING, text,
            key=f"h{level}-{slug(text) or index}", locator=locator,
            metadata={"level": level, "style": style})
        if block is not None:
            builder.push_heading(level, block.block_key, text)
        return

    locator["headingPath"] = list(builder.heading_path)
    if _CODE_STYLE_RE.search(style):
        block_type = DocumentBlockType.CODE_BLOCK
    elif _LIST_STYLE_RE.search(style):
        block_type = DocumentBlockType.LIST_ITEM
    else:
        block_type = DocumentBlockType.PARAGRAPH
    builder.add(block_type, text, key=f"p-{index:06d}", locator=locator,
                metadata={"style": style} if style else {})


def _emit_table(builder: DocumentBuilder, key: str, table: Any,
                table_index: int) -> None:
    rows = list(table.rows)
    header: List[str] = []
    if rows:
        header = [(_cell_text(c)) for c in rows[0].cells]

    table_key = f"table-{table_index:04d}"
    builder.add(
        DocumentBlockType.TABLE, " | ".join(header), key=table_key,
        locator={"kind": "docx-table-row", "document": key,
                 "tableIndex": table_index, "rowIndex": 0,
                 "cellRange": f"A1:{_column_letter(max(1, len(header)))}1"},
        content_mode=ContentMode.DERIVED, indexable=False,
        metadata={"tableIndex": table_index, "rows": len(rows),
                  "columns": len(header), "headers": ", ".join(header)})

    for row_no, row in enumerate(rows, start=1):
        if builder.full:
            break
        cells = [_cell_text(c) for c in row.cells]
        if not any(c.strip() for c in cells):
            continue
        # The header row is emitted too (it is real content someone may search
        # for), but labelled positionally — pairing it with itself would produce
        # the useless "Requirement: Requirement".
        names = ([_column_letter(i + 1) for i in range(len(cells))]
                 if row_no == 1 else header)
        pairs = [f"{names[i] if i < len(names) and names[i] else f'col{i + 1}'}"
                 f": {value}"
                 for i, value in enumerate(cells) if value.strip()]
        builder.add(
            DocumentBlockType.TABLE_ROW, "; ".join(pairs),
            key=f"{table_key}-r{row_no:04d}",
            locator={"kind": "docx-table-row", "document": key,
                     "tableIndex": table_index, "rowIndex": row_no,
                     "cellRange": f"A{row_no}:"
                                  f"{_column_letter(max(1, len(cells)))}{row_no}"},
            # DERIVED: `header: value` is a rendering of the row, not the
            # document's literal text. The locator still cites the exact row.
            content_mode=ContentMode.DERIVED, parent_key=table_key,
            metadata={"row": row_no, "cells": len(cells),
                      "header_row": row_no == 1})


def _cell_text(cell: Any) -> str:
    try:
        return " ".join((cell.text or "").split())
    except Exception:
        return ""


__all__ = ["DocxParser", "CLAIMS"]
